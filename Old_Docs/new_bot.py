import os
import re
import datetime
import pypandoc
from linkedin_api import Linkedin
import telebot
from PIL import Image
from bs4 import BeautifulSoup
import docx
from docx.shared import Cm
import requests
from matplotlib import pyplot as plt


with open('/Users/monkey/Public/Python/Hidden files/LinkedIn_Bot.txt', 'r') as file:
    key = file.readline().rstrip('\n')
    mail = file.readline().rstrip('\n')
    passwd = file.readline().rstrip('\n')
bot = telebot.TeleBot(key)
# Check if bot is connected
try:
    bot_info = bot.get_me()
    print('Bot is connected:', bot_info.first_name)
except Exception as e:
    print('Bot is not connected:', e)

# Define your button labels
start_button_label = "Start dossier"
png_button_label = "PNG"
pdf_button_label = "PDF"
menu_button_label = "Menu"
next_button_label = "Next"
# Define your button payloads
start_btn_pl = "start_press"
png_btn_pl = "png_press"
pdf_btn_pl = "pdf_press"
menu_btn_pl = "menu_press"
next_btn_pl = "next_press"

# Define buttons
start_press = telebot.types.InlineKeyboardButton(start_button_label, callback_data=start_btn_pl)
png_press = telebot.types.InlineKeyboardButton(png_button_label, callback_data=png_btn_pl)
pdf_press = telebot.types.InlineKeyboardButton(pdf_button_label, callback_data=pdf_btn_pl)
menu_press = telebot.types.InlineKeyboardButton(menu_button_label, callback_data=menu_btn_pl)
next_press = telebot.types.InlineKeyboardButton(next_button_label, callback_data=next_btn_pl)

# Define your 1 keyboard
keyboard1 = telebot.types.InlineKeyboardMarkup(row_width=1)
keyboard1.add(start_press)
# Define your 2 keyboard
keyboard2 = telebot.types.InlineKeyboardMarkup(row_width=1)
keyboard2.add(menu_press)
# Define your 3 keyboard
keyboard3 = telebot.types.InlineKeyboardMarkup(row_width=2)
keyboard3.add(png_press, menu_press)
# Define your 4 keyboard
keyboard4 = telebot.types.InlineKeyboardMarkup(row_width=1)
keyboard4.add(next_press)


# Send a message with the first keyboard to the user
@bot.message_handler(commands=['start'])
def send_welcome(message):
    bot.send_message(message.chat.id, "Welcome to Stat Bot! Choose the site:", reply_markup=keyboard1)


# Handle button taps
@bot.callback_query_handler(func=lambda call: True)
def handle_button_tap(call):
    if call.data == start_btn_pl:
        bot.send_message(call.message.chat.id, "Please enter a valid GitHub link:", reply_markup=keyboard2)
        bot.register_next_step_handler(call.message, get_github_link)
    elif call.data == next_btn_pl:
        bot.send_message(call.message.chat.id, "First part done! \nPlease enter a valid LinkedIn link:",
                         reply_markup=keyboard2)
        bot.register_next_step_handler(call.message, get_linkedin_link)
    elif call.data == png_btn_pl:
        png_gen(call.message.chat.id)
    elif call.data == pdf_btn_pl:
        pdf_gen(call.message.chat.id)
    elif call.data == menu_btn_pl:
        bot.send_message(call.message.chat.id, "\tMenu!\n Choose the site:", reply_markup=keyboard1)
    else:
        bot.send_message(call.message.chat.id, "Unknown button tapped.")
    bot.edit_message_reply_markup(chat_id=call.message.chat.id, message_id=call.message.message_id, reply_markup=None)


def get_github_link(message):
    github_link = message.text
    if github_link.startswith("https://github.com/") or github_link.startswith("github.com/"):
        # Open the Word document
        doc = docx.Document('/Users/monkey/Public/Python/Python_Bot/template.docx')
        # Replace the link with the URL of the image you want to download
        url = github_link
        response = requests.get(url)
        # Parse the HTML content of the response using BeautifulSoup
        soup = BeautifulSoup(response.content, 'html.parser')
        # Extract the name of the GitHub user or organization
        name_elem = soup.find('span', class_='p-name vcard-fullname d-block overflow-hidden')
        username = name_elem.text.strip() if name_elem else None
        # Find the element that contains the number of followers and extract the text
        followers_elem = soup.find('a', {'href': f'/{username}?tab=followers'})
        followers_num = followers_elem.find('span', {'class': 'text-bold'}).text.strip() if followers_elem else None
        # Extract the number of repositories
        num_repos_elem = soup.find('span', class_='Counter')
        repos_num = num_repos_elem.text.strip() if num_repos_elem else None

        url = f"https://api.github.com/users/{username}"
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            image_url = data["avatar_url"]
            # Make a GET request to the image URL and save the image to a file
            image_response = requests.get(image_url)
            with open(f"/Users/monkey/Public/Python/Python_Bot/Users_docs/profile_img_{message.chat.id}.png", "wb") as f:
                f.write(image_response.content)
        else:
            print(f"Failed to download profile image for {username}. Status code: {response.status_code}")

        # Insert the image at the center of the first line
        first_paragraph = doc.paragraphs[0]
        run = first_paragraph.add_run()
        # Open the image file
        img = Image.open(f'/Users/monkey/Public/Python/Python_Bot/Users_docs/profile_img_{message.chat.id}.png')
        # Convert the image to PNG format
        img = img.convert('RGB')
        # Save the image as a PNG file
        img.save(f'/Users/monkey/Public/Python/Python_Bot/Users_docs/profile_img_{message.chat.id}.png', 'PNG')
        run.add_picture(f'/Users/monkey/Public/Python/Python_Bot/Users_docs/profile_img_{message.chat.id}.png', width=Cm(4), height=Cm(4))
        run.alignment = 1  # Set the alignment to center

        # Loop through all paragraphs in the document
        for paragraph in doc.paragraphs:
            if '<username>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<username>', username)
            if '<repos_num>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<repos_num>', repos_num)
            if '<followers_num>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<followers_num>', "followers_num")

        # Insert the image at the end of the last line
        last_paragraph = doc.paragraphs[-1]
        run = last_paragraph.add_run()

        import subprocess

        # Run the 'my_script.py' file with 'my_var' as a command-line argument
        subprocess.run(['python', 'add_stat.py', username])

        run.add_picture(f'Im_Stat_{username}.png', width=Cm(12), height=Cm(6))
        run.add_break()
        run.alignment = 1  # Set the alignment to center

        # Save the modified document
        doc.save(f'/Users/monkey/Public/Python/Python_Bot/Users_docs/dossier_{message.chat.id}.docx')

        bot.send_message(message.chat.id, "Done, click NEXT button!", reply_markup=keyboard4)

    else:
        bot.send_message(message.chat.id,
                         "Please enter a valid GitHub link starting with 'https://github.com/' or 'github.com/'",
                         reply_markup=keyboard2)
        bot.register_next_step_handler(message, get_github_link)


# Linkedin part
def get_linkedin_link(message):
    linkedin_link = message.text
    if linkedin_link.startswith("https://www.linkedin.com/") or linkedin_link.startswith("www.linkedin.com/"):

        # Authenticate using bot account
        api = Linkedin(mail, passwd)
        parts = linkedin_link.split("/")
        username = parts[-1]
        # GET a profile
        profile = api.get_profile(username)

        first_name = profile['firstName']
        last_name = profile['lastName']
        employment = profile['headline']

        location = profile['locationName'] + profile['geoLocationName']
        bio = profile['summary']

        # Open the Word document
        doc = docx.Document(f'/Users/monkey/Public/Python/Python_Bot/Users_docs/dossier_{message.chat.id}.docx')

        # Writing data
        for paragraph in doc.paragraphs:
            if '<first_name>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<first_name>', first_name)
            if '<last_name>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<last_name>', last_name)
            if '<employment>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<employment>', employment)
            if '<bio>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<bio>', bio)
            if '<followers_num>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<followers_num>', "followers_num")
            if '<location>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<location>', location)
            if '<skills>' in paragraph.text:
                skills_list = profile['skills']
                skills_str = ", ".join(skill['name'] for skill in skills_list)
                paragraph.text = paragraph.text.replace('<skills>', skills_str)
            if '<education>' in paragraph.text:
                education_list = []
                for edu in profile['education']:
                    if 'school' in edu:
                        school_name = edu['school']['schoolName']
                        education_list.append(school_name)
                    else:
                        print("No school information available")
                # join the list of school names with commas
                education_str = ", ".join(education_list)
                # replace <education> in paragraph.text with the comma-separated string
                paragraph.text = paragraph.text.replace('<education>', education_str)

        # Save the modified document
        doc.save(f'/Users/monkey/Public/Python/Python_Bot/Users_docs/dossier_{message.chat.id}.docx')

        # Sending a file
        bot.send_document(chat_id=message.chat.id,
                          document=open(
                              f'/Users/monkey/Public/Python/Python_Bot/Users_docs/dossier_{message.chat.id}.docx',
                              'rb'))
        os.system("touch " + "/Users/monkey/Public/Python/Python_Bot/Users_docs")
        bot.send_message(message.chat.id, "Start new :)", reply_markup=keyboard1)

    else:
        bot.send_message(message.chat.id, "Please enter a valid LinkedIn link starting with 'https://www.linkedin.com/'"
                                          " or 'www.linkedin.com/'", reply_markup=keyboard2)
        bot.register_next_step_handler(message, get_linkedin_link)


def png_gen(chat_id):

    bot.send_message(chat_id, "Try another link:", reply_markup=keyboard3)


def pdf_gen(chat_id):
    # Path of the input Word document
    input_file_path = f'/Users/monkey/Public/Python/Python_Bot/Users_docs/dossier_{chat_id}.docx'
    output_file_path = f'/Users/monkey/Public/Python/Python_Bot/Users_docs/dossier_{chat_id}.pdf'

    output = pypandoc.convert_file(input_file_path, "pdf", outputfile=output_file_path)

    # Sending a file
    bot.send_document(chat_id,
                      document=open(f'/Users/monkey/Public/Python/Python_Bot/Users_docs/dossier_{chat_id}.pdf', 'rb'))
    bot.send_message(chat_id, "Try another link:", reply_markup=keyboard2)


def remove_non_ascii(text):
    return re.sub(r'[^\x00-\x7F]+', '', text)


def del_dir():
    directory = "/Users/monkey/Public/Python/Python_Bot/Users_docs"
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        try:
            if os.path.isfile(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"Error deleting file: {file_path}. {e}")


# del_dir()
# Start the bot
bot.polling()
