from github import Github
from funcs import get_all_commits_by_user
from linkedin_api import Linkedin
import telebot
from PIL import Image
import docx
from docx.shared import Cm
import requests
import os
import re

# Read bot credentials from file
with open('/Users/monkey/Public/Python/Hidden files/LinkedIn_Bot.txt', 'r') as file:
    key = file.readline().rstrip('\n')
    mail = file.readline().rstrip('\n')
    passwd = file.readline().rstrip('\n')
    git_tok = file.readline().rstrip('\n')

# Initialize telebot
bot = telebot.TeleBot(key)

# Define button labels and payloads
start_button_label = "Start dossier"
menu_button_label = "Menu"
next_button_label = "Next"
start_btn_pl = "start_press"
menu_btn_pl = "menu_press"
next_btn_pl = "next_press"

# Define buttons
start_press = telebot.types.InlineKeyboardButton(start_button_label, callback_data=start_btn_pl)
menu_press = telebot.types.InlineKeyboardButton(menu_button_label, callback_data=menu_btn_pl)
next_press = telebot.types.InlineKeyboardButton(next_button_label, callback_data=next_btn_pl)

# Define keyboards
keyboard1 = telebot.types.InlineKeyboardMarkup(row_width=1)
keyboard1.add(start_press)
keyboard2 = telebot.types.InlineKeyboardMarkup(row_width=1)
keyboard2.add(menu_press)
keyboard4 = telebot.types.InlineKeyboardMarkup(row_width=1)
keyboard4.add(next_press)


# Send welcome message with keyboard1
@bot.message_handler(commands=['start'])
def send_welcome(message):
    bot.send_message(message.chat.id, "Welcome to Stat Bot! Choose the site:", reply_markup=keyboard1)


# Handle button taps
@bot.callback_query_handler(func=lambda call: True)
def handle_button_tap(call):
    if call.data == start_btn_pl:
        bot.send_message(call.message.chat.id, "Please enter a valid GitHub link:", reply_markup=keyboard2)
        bot.register_next_step_handler(call.message, get_github_link)
    elif call.data == next_btn_pl:
        bot.send_message(call.message.chat.id, "First part done! "
                                               "\nPlease enter a valid LinkedIn link:", reply_markup=keyboard2)
        bot.register_next_step_handler(call.message, get_linkedin_link)
    elif call.data == menu_btn_pl:
        bot.send_message(call.message.chat.id, "\tMenu!\n Choose the site:", reply_markup=keyboard1)
    else:
        bot.send_message(call.message.chat.id, "Unknown button tapped.")
    bot.edit_message_reply_markup(chat_id=call.message.chat.id, message_id=call.message.message_id, reply_markup=None)


def get_github_link(message):
    if message.text.startswith("https://github.com/") or message.text.startswith("github.com/"):
        # Call Func
        gitHubFill(message)
        bot.send_message(message.chat.id, "Done, click NEXT button!", reply_markup=keyboard4)
    else:
        bot.send_message(message.chat.id, "Please enter a valid GitHub link starting "
                                          "with 'https://github.com/' or 'github.com/'", reply_markup=keyboard2)
        bot.register_next_step_handler(message, get_github_link)


def gitHubFill(message):
    url = message.text
    username = url[len('https://github.com/'):].split('/')[0]
    print(username)
    g = Github(git_tok)
    # Func About Commits
    get_all_commits_by_user(username)

    user = g.get_user(username)

    # Get user information
    repos_num = user.public_repos
    num_stars = user.get_starred().totalCount
    num_followers = user.followers
    num_following = user.following

    # Get profile image
    image_url = user.avatar_url
    response = requests.get(image_url)
    with open(f"Users_docs/profile_img_{message.chat.id}.png", "wb") as f:
        f.write(response.content)

    # Insert the image at the center of the first line
    doc = docx.Document('Old_Docs/template.docx')
    first_paragraph = doc.paragraphs[0]
    run = first_paragraph.add_run()
    img = Image.open(f'Users_docs/profile_img_{message.chat.id}.png')
    img = img.convert('RGB')
    img.save(f'Users_docs/profile_img_{message.chat.id}.png', 'PNG')
    run.add_picture(f'Users_docs/profile_img_{message.chat.id}.png', width=Cm(4),
                    height=Cm(4))
    run.alignment = 1  # Set the alignment to center

    # Loop through all paragraphs in the document
    for paragraph in doc.paragraphs:
        if '<username>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<username>', username)
        if '<repos_num>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<repos_num>', str(repos_num))
        if '<stars_num>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<stars_num>', str(num_stars))
        if '<followers_num>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<followers_num>', str(num_followers))
        if '<following_num>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<following_num>', str(num_following))

    # Insert the image at the end of the last line
    last_paragraph = doc.paragraphs[-1]
    run = last_paragraph.add_run()

    run.add_picture(f'Users_docs/Im_Stat_{username}.png', width=Cm(18), height=Cm(8))
    run.add_break()
    run.alignment = 1  # Set the alignment to center

    # Save the modified document
    doc.save(f'Users_docs/dossier_{message.chat.id}.docx')


def get_linkedin_link(message):
    linkedin_link = message.text
    if linkedin_link.startswith("https://www.linkedin.com/") or linkedin_link.startswith("www.linkedin.com/"):
        # Authenticate using bot account
        api = Linkedin(mail, passwd)
        username_pattern = r"in\/([^\?\/]*)"
        match = re.search(username_pattern, linkedin_link)
        if match:
            username = match.group(1)
        else:
            username = None

        # GET a profile
        profile = api.get_profile(username)

        first_name = profile['firstName']
        last_name = profile['lastName']
        employment = profile['headline']

        location = profile['locationName'] + profile['geoLocationName']
        if profile['summary']:
            bio = profile['summary']
        else:
            bio = None

        # Open the Word document
        doc = docx.Document(f'Users_docs/dossier_{message.chat.id}.docx')

        # Writing data
        for paragraph in doc.paragraphs:
            if '<first_name>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<first_name>', first_name)
            if '<last_name>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<last_name>', last_name)
            if '<employment>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<employment>', employment)
            if '<bio>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<bio>', bio)
            if '<followers_num>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<followers_num>', "followers_num")
            if '<location>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<location>', location)
            if '<skills>' in paragraph.text:
                skills_list = profile['skills']
                skills_str = ", ".join(skill['name'] for skill in skills_list)
                paragraph.text = paragraph.text.replace('<skills>', skills_str)
            if '<education>' in paragraph.text:
                education_list = []
                for edu in profile['education']:
                    if 'school' in edu:
                        school_name = edu['school']['schoolName']
                        education_list.append(school_name)
                    else:
                        print("No school information available")
                # join the list of school names with commas
                education_str = ", ".join(education_list)
                # replace <education> in paragraph.text with the comma-separated string
                paragraph.text = paragraph.text.replace('<education>', education_str)

        # Save the modified document
        doc.save(f'Users_docs/dossier_{message.chat.id}.docx')

        # Sending a file
        bot.send_document(chat_id=message.chat.id,
                          document=open(
                              f'/Users/monkey/Public/Python/Python_Bot/Users_docs/dossier_{message.chat.id}.docx',
                              'rb'))
        os.system("touch " + "/Users/monkey/Public/Python/Python_Bot/Users_docs")
        bot.send_message(message.chat.id, "Start new :)", reply_markup=keyboard1)

    else:
        bot.send_message(message.chat.id, "Please enter a valid LinkedIn link starting with 'https://www.linkedin.com/'"
                                          " or 'www.linkedin.com/'", reply_markup=keyboard2)
        bot.register_next_step_handler(message, get_linkedin_link)


if __name__ == '__main__':
    bot.polling()
